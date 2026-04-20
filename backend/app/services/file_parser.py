"""文件解析服务 (TDD §3.1 · 文件流解析)。

策略：
1. 不信任文件扩展名 —— 先用 `filetype` 做 magic number 嗅探；
2. PDF 走 PyMuPDF(fitz)，DOCX 走 python-docx，纯文本 utf-8/gbk 兜底；
3. 对图片 / 加密 PDF / 未知类型返回 `UNSUPPORTED`，由上层决定降级策略。
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from typing import Literal

import filetype

from app.models.file import ParseStatus

logger = logging.getLogger(__name__)


MimeKind = Literal[
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
    "image/jpeg",
    "image/png",
    "unknown",
]


@dataclass(slots=True)
class ParseResult:
    """文件解析结果。"""

    status: ParseStatus
    mime: str
    text: str
    error: str | None = None


# 扩展名作为 magic 嗅探失败时的回退。ZIP 压缩包里 Boss 直聘的乱码文件名常常
# 丢扩展名，此时只依赖 content 嗅探。
_EXT_TO_MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
    "txt": "text/plain",
    "md": "text/plain",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
}


def detect_mime(data: bytes, filename: str = "") -> str:
    """优先 magic number 嗅探，失败再退回扩展名。"""
    kind = filetype.guess(data)
    if kind is not None:
        return kind.mime
    if "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()
        if ext in _EXT_TO_MIME:
            return _EXT_TO_MIME[ext]
    # 启发式：看起来像 UTF-8 文本
    try:
        data[:4096].decode("utf-8")
        return "text/plain"
    except UnicodeDecodeError:
        return "unknown"


def _parse_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF, 延迟导入避免冷启动开销

    buf: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        if getattr(doc, "needs_pass", False):
            raise ValueError("PDF 已加密，无法提取文本")
        for page in doc:
            buf.append(page.get_text("text"))
    text = "\n".join(buf).strip()
    if not text:
        raise ValueError("PDF 无可提取文本（可能是扫描件/图片）")
    return text


def _parse_docx(data: bytes) -> str:
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("DOCX 正文为空")
    return text


def _parse_plain(data: bytes) -> str:
    for encoding in ("utf-8", "gbk", "gb18030", "utf-16"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("无法按常见中文编码解码纯文本")


def parse_file(data: bytes, filename: str = "") -> ParseResult:
    """解析单个文件为文本。"""
    mime = detect_mime(data, filename)
    try:
        if mime == "application/pdf":
            return ParseResult(ParseStatus.PARSED, mime, _parse_pdf(data))
        if mime.endswith("wordprocessingml.document"):
            return ParseResult(ParseStatus.PARSED, mime, _parse_docx(data))
        if mime == "application/msword":
            # 旧版 .doc 格式 python-docx 不支持，直接标记不支持
            return ParseResult(
                ParseStatus.UNSUPPORTED,
                mime,
                "",
                "旧版 .doc 格式需手动另存为 .docx",
            )
        if mime.startswith("text/"):
            return ParseResult(ParseStatus.PARSED, mime, _parse_plain(data))
        if mime.startswith("image/"):
            return ParseResult(
                ParseStatus.UNSUPPORTED,
                mime,
                "",
                "图片需 OCR，Phase 3 暂不处理",
            )
        return ParseResult(
            ParseStatus.UNSUPPORTED, mime, "", f"暂不支持的 MIME：{mime}"
        )
    except Exception as exc:  # noqa: BLE001 —— 统一降级为 FAILED
        logger.warning("parse_file failed mime=%s err=%s", mime, exc)
        return ParseResult(ParseStatus.FAILED, mime, "", str(exc))
